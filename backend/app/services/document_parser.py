"""
Document parser service for converting Word and Markdown files to PRD format
"""

import logging
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document
from app.models.prd_models import PRD, PRDSection, Priority
import uuid

logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """Parse Word documents and Markdown files into PRD structure"""

    # Common section patterns that indicate PRD sections
    SECTION_PATTERNS = {
        'overview': ['overview', 'summary', 'introduction', 'executive summary'],
        'objectives': ['objectives', 'goals', 'purpose'],
        'requirements': ['requirements', 'functional requirements', 'specifications'],
        'features': ['features', 'capabilities', 'functionality'],
        'constraints': ['constraints', 'limitations', 'restrictions'],
        'stakeholders': ['stakeholders', 'users', 'audience'],
        'metrics': ['metrics', 'kpis', 'success criteria', 'measurements'],
        'dependencies': ['dependencies', 'prerequisites', 'integrations'],
        'risks': ['risks', 'challenges', 'concerns'],
        'timeline': ['timeline', 'schedule', 'milestones'],
    }

    @staticmethod
    def _determine_section_priority(title: str, content: str) -> Priority:
        """
        Determine priority based on section title and content keywords

        Args:
            title: Section title
            content: Section content

        Returns:
            Priority enum value
        """
        title_lower = title.lower()
        content_lower = content.lower()

        # Critical indicators
        critical_keywords = ['critical', 'must have', 'essential', 'required', 'mandatory']
        if any(keyword in title_lower or keyword in content_lower for keyword in critical_keywords):
            return Priority.CRITICAL

        # High priority indicators
        high_keywords = ['important', 'high priority', 'key', 'core']
        if any(keyword in title_lower or keyword in content_lower for keyword in high_keywords):
            return Priority.HIGH

        # Low priority indicators
        low_keywords = ['nice to have', 'optional', 'future', 'low priority']
        if any(keyword in title_lower or keyword in content_lower for keyword in low_keywords):
            return Priority.LOW

        # Default to medium
        return Priority.MEDIUM

    @staticmethod
    def _extract_tags(title: str, content: str) -> List[str]:
        """
        Extract tags from section title and content

        Args:
            title: Section title
            content: Section content

        Returns:
            List of tags
        """
        tags = []
        title_lower = title.lower()

        # Add tags based on section type
        for tag, patterns in DocumentParser.SECTION_PATTERNS.items():
            if any(pattern in title_lower for pattern in patterns):
                tags.append(tag)

        # Extract hashtags from content (if present)
        hashtags = re.findall(r'#(\w+)', content)
        tags.extend(hashtags)

        return list(set(tags))  # Remove duplicates

    @staticmethod
    def parse_docx(file_path: str, prd_name: Optional[str] = None, prd_description: Optional[str] = None) -> PRD:
        """
        Parse a Word document (.docx) into a PRD

        Args:
            file_path: Path to the .docx file
            prd_name: Optional name for the PRD (defaults to filename)
            prd_description: Optional description

        Returns:
            PRD object

        Raises:
            DocumentParserError: If parsing fails
        """
        try:
            doc = Document(file_path)
            path_obj = Path(file_path)

            # Use provided name or filename (without extension)
            name = prd_name or path_obj.stem

            sections: List[PRDSection] = []
            current_section_title = None
            current_section_content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a heading (assuming style contains 'Heading')
                is_heading = 'Heading' in para.style.name

                if is_heading and current_section_title:
                    # Save the previous section
                    content = '\n'.join(current_section_content)
                    if content:
                        priority = DocumentParser._determine_section_priority(
                            current_section_title, content
                        )
                        tags = DocumentParser._extract_tags(current_section_title, content)
                        sections.append(PRDSection(
                            title=current_section_title,
                            content=content,
                            priority=priority,
                            tags=tags
                        ))

                    # Start new section
                    current_section_title = text
                    current_section_content = []

                elif is_heading:
                    # First heading encountered
                    current_section_title = text
                    current_section_content = []

                else:
                    # Regular paragraph - add to current section
                    if current_section_title:
                        current_section_content.append(text)
                    else:
                        # Content before any heading - create an "Overview" section
                        if not any(s.title == "Overview" for s in sections):
                            current_section_title = "Overview"
                            current_section_content = [text]

            # Don't forget the last section
            if current_section_title and current_section_content:
                content = '\n'.join(current_section_content)
                priority = DocumentParser._determine_section_priority(
                    current_section_title, content
                )
                tags = DocumentParser._extract_tags(current_section_title, content)
                sections.append(PRDSection(
                    title=current_section_title,
                    content=content,
                    priority=priority,
                    tags=tags
                ))

            # If no sections were found, create a single section with all text
            if not sections:
                all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                if all_text:
                    sections.append(PRDSection(
                        title="Document Content",
                        content=all_text,
                        priority=Priority.MEDIUM,
                        tags=[]
                    ))

            # Extract description from first section if not provided
            if not prd_description and sections:
                # Use first 200 chars of first section as description
                prd_description = sections[0].content[:200] + "..." if len(sections[0].content) > 200 else sections[0].content

            prd = PRD(
                id=str(uuid.uuid4()),
                name=name,
                description=prd_description,
                sections=sections
            )

            logger.info(f"Successfully parsed Word document: {file_path}")
            logger.info(f"Created PRD with {len(sections)} sections")

            return prd

        except Exception as e:
            raise DocumentParserError(f"Failed to parse Word document {file_path}: {str(e)}")

    @staticmethod
    def parse_markdown(file_path: str, prd_name: Optional[str] = None, prd_description: Optional[str] = None) -> PRD:
        """
        Parse a Markdown file into a PRD

        Args:
            file_path: Path to the .md file
            prd_name: Optional name for the PRD (defaults to filename)
            prd_description: Optional description

        Returns:
            PRD object

        Raises:
            DocumentParserError: If parsing fails
        """
        try:
            path_obj = Path(file_path)

            # Use provided name or filename (without extension)
            name = prd_name or path_obj.stem

            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            sections: List[PRDSection] = []

            # Split by headings (# or ##)
            # Pattern matches markdown headings
            heading_pattern = r'^(#{1,6})\s+(.+)$'

            lines = content.split('\n')
            current_section_title = None
            current_section_content = []

            for line in lines:
                match = re.match(heading_pattern, line)

                if match:
                    # Save previous section
                    if current_section_title:
                        section_content = '\n'.join(current_section_content).strip()
                        if section_content:
                            priority = DocumentParser._determine_section_priority(
                                current_section_title, section_content
                            )
                            tags = DocumentParser._extract_tags(current_section_title, section_content)
                            sections.append(PRDSection(
                                title=current_section_title,
                                content=section_content,
                                priority=priority,
                                tags=tags
                            ))

                    # Start new section
                    heading_level = len(match.group(1))
                    current_section_title = match.group(2).strip()
                    current_section_content = []
                else:
                    # Regular line - add to current section
                    if current_section_title is not None:
                        current_section_content.append(line)
                    elif line.strip():
                        # Content before any heading
                        if not any(s.title == "Overview" for s in sections):
                            current_section_title = "Overview"
                            current_section_content = [line]

            # Don't forget the last section
            if current_section_title and current_section_content:
                section_content = '\n'.join(current_section_content).strip()
                if section_content:
                    priority = DocumentParser._determine_section_priority(
                        current_section_title, section_content
                    )
                    tags = DocumentParser._extract_tags(current_section_title, section_content)
                    sections.append(PRDSection(
                        title=current_section_title,
                        content=section_content,
                        priority=priority,
                        tags=tags
                    ))

            # If no sections were found, create a single section with all content
            if not sections:
                if content.strip():
                    sections.append(PRDSection(
                        title="Document Content",
                        content=content.strip(),
                        priority=Priority.MEDIUM,
                        tags=[]
                    ))

            # Extract description from first section if not provided
            if not prd_description and sections:
                # Use first 200 chars of first section as description
                prd_description = sections[0].content[:200] + "..." if len(sections[0].content) > 200 else sections[0].content

            prd = PRD(
                id=str(uuid.uuid4()),
                name=name,
                description=prd_description,
                sections=sections
            )

            logger.info(f"Successfully parsed Markdown file: {file_path}")
            logger.info(f"Created PRD with {len(sections)} sections")

            return prd

        except Exception as e:
            raise DocumentParserError(f"Failed to parse Markdown file {file_path}: {str(e)}")

    @staticmethod
    def parse_document(file_path: str, prd_name: Optional[str] = None, prd_description: Optional[str] = None) -> PRD:
        """
        Auto-detect file type and parse into PRD

        Args:
            file_path: Path to the document file
            prd_name: Optional name for the PRD
            prd_description: Optional description

        Returns:
            PRD object

        Raises:
            DocumentParserError: If file type is unsupported or parsing fails
        """
        path_obj = Path(file_path)

        if not path_obj.exists():
            raise DocumentParserError(f"File not found: {file_path}")

        suffix = path_obj.suffix.lower()

        if suffix == '.docx':
            return DocumentParser.parse_docx(file_path, prd_name, prd_description)
        elif suffix in ['.md', '.markdown']:
            return DocumentParser.parse_markdown(file_path, prd_name, prd_description)
        else:
            raise DocumentParserError(f"Unsupported file type: {suffix}. Supported types: .docx, .md, .markdown")
